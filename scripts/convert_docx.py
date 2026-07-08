#!/usr/bin/env python3
"""
Convert DOCX to Knot HTML with proper heading hierarchy, images, and tables.

Heading rules:
  H1 — bold paragraphs starting with Roman numerals (I., II., III.)
  H2 — bold paragraphs with indent < 3 characters (major sections)
  H3 — bold paragraphs with indent >= 3 characters, OR bold "Bước X:" after H2
  H4 — "Bước X:" paragraphs (bold or not) appearing when previous heading was H3
"""

import docx
from docx.oxml.ns import qn
import base64
import io
import sqlite3
import re
from PIL import Image

DOCX_PATH = "/Users/chad/Documents/Tai lieu TH ĐTTT-16.9.2023.docx"
DB_PATH = "/Users/chad/Documents/AI/knot/data/knot.db"

doc = docx.Document(DOCX_PATH)
body = doc.element.body
children = list(body)

NOTE_IDS = {
    1: "e513a14e811e485a8ae1",
    2: "af066bda863749e3b832",
    3: "ead801aa08ab4a81974b",
    4: "1f2c7d2541e8487cb8ba",
    5: "89f244a7b5d240278ebb",
}

BAI_NAMES = {
    1: "Bài 1 - Hướng dẫn sử dụng KIT thực hành NI ELVIS II/II+",
    2: "Bài 2 - Khảo sát đáp ứng tần số của mạch khuếch đại",
    3: "Bài 3 - Khảo sát mạch trễ pha",
    4: "Bài 4 - Khảo sát mạch lọc tích cực",
    5: "Bài 5 - Khảo sát các khối điện tử thông tin",
}

# --- helpers ---

def get_table_text(t):
    return ' '.join(c.text for row in t.rows for c in row.cells)

def extract_img_data(run_elem):
    for drawing in run_elem.findall('.//' + qn('w:drawing')):
        for blip in drawing.findall('.//' + qn('a:blip')):
            embed = blip.get(qn('r:embed'))
            if embed and embed in doc.part.rels:
                rel = doc.part.rels[embed]
                img_data = rel.target_part.blob
                img = Image.open(io.BytesIO(img_data))
                b64 = base64.b64encode(img_data).decode()
                return f"data:image/{img.format.lower()};base64,{b64}"
    return None

def is_bold(p):
    return any(r.bold for r in p.runs if r.bold)

def get_indent_chars(p):
    """Get left indent in character units (1ch ≈ 240 twips)."""
    pPr = p._element.find(qn('w:pPr'))
    if pPr is not None:
        ind = pPr.find(qn('w:ind'))
        if ind is not None:
            left = ind.get(qn('w:left'))
            if left:
                return int(left) // 240
    return 0

def escape(s):
    return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace('"', "&quot;")

RX_ROMAN = re.compile(r'^[IVXLCDM]+\.\s')
RX_FIGURE = re.compile(r'^Hình\s+(\d+\.\d+)[:\s.]+(.+)')
RX_BuOC = re.compile(r'^(Bước|bước|BUỚC)\s+\d+[:]')
RX_NOTE = re.compile(r'^(Lưu ý|Chú ý|Ghi chú)\s*(\d+\s*)?[:]', re.IGNORECASE)
RX_YEU_CAU = re.compile(r'^Yêu cầu\s+\d+\s*[:]', re.IGNORECASE)
RX_INSTRUCTION_START = re.compile(
        r'^(Để\s+(lấy|thay|tạo|thực|có|được)|'
        r'Tiến hành|Thực hiện|Nhấn|Click|'
        r'Vẽ\s+(sơ đồ|mạch)|Chọn|'
        r'Tại\s+giao\s+diện|Sử\s+dụng|Dùng|'
        r'Lưu\s+ý\s*[:]|'
        r'Truy cập)',
    re.IGNORECASE
)

# --- map bài boundaries ---
bai_markers = {}
table_idx = 0
para_idx = 0
for idx, elem in enumerate(children):
    tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
    if tag == 'tbl':
        t = doc.tables[table_idx]
        txt = get_table_text(t)
        for b in range(1, 6):
            if f'BÀI {b}' in txt and b not in bai_markers.values():
                bai_markers[idx] = b
                break
        table_idx += 1
    elif tag == 'p':
        para_idx += 1

sorted_markers = sorted(bai_markers.items())
bai_ranges = {}
for i, (idx, b) in enumerate(sorted_markers):
    end_idx = sorted_markers[i+1][0] if i+1 < len(sorted_markers) else len(children)
    bai_ranges[b] = (idx, end_idx)

# --- pre-extract all items with metadata ---
table_idx = 0
para_idx = 0
all_items = []

for idx, elem in enumerate(children):
    tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
    if tag == 'tbl':
        t = doc.tables[table_idx]
        txt = get_table_text(t)
        bai = 0
        for b, (start, end) in bai_ranges.items():
            if start <= idx < end:
                bai = b
                break
        all_items.append({'type': 'table', 'idx': idx, 'table': t, 'text': txt, 'bai': bai})
        table_idx += 1
    elif tag == 'p':
        p = doc.paragraphs[para_idx]
        text = p.text.strip()
        bai = 0
        for b, (start, end) in bai_ranges.items():
            if start <= idx < end:
                bai = b
                break
        imgs = []
        for run_elem in p._element.findall(qn('w:r')):
            b64 = extract_img_data(run_elem)
            if b64:
                imgs.append(b64)
        all_items.append({
            'type': 'p', 'idx': idx, 'p': p, 'text': text,
            'bold': is_bold(p), 'imgs': imgs, 'indent': get_indent_chars(p),
            'bai': bai,
        })
        para_idx += 1

# --- table to HTML ---
def table_to_html(t):
    rows_html = []
    for row in t.rows:
        # Use raw <w:tc> elements to avoid python-docx virtual cells
        tc_list = row._tr.findall(qn('w:tc'))
        cells_html = []
        for tc in tc_list:
            tcPr = tc.find(qn('w:tcPr'))
            colspan = 1
            if tcPr is not None:
                gs = tcPr.find(qn('w:gridSpan'))
                if gs is not None:
                    colspan = int(gs.get(qn('w:val')) or 1)
            parts = []
            for p_elem in tc.findall(qn('w:p')):
                ptext = ''.join(r.text or '' for r in p_elem.findall('.//' + qn('w:t')))
                if ptext.strip():
                    parts.append(f'<p>{escape(ptext.strip())}</p>')
            cells_html.append(f'<td colspan="{colspan}" rowspan="1" class="">{"".join(parts)}</td>')
        rows_html.append('<tr>' + ''.join(cells_html) + '</tr>')
    tblGrid = t._tbl.find(qn('w:tblGrid'))
    grid_cols = len(tblGrid.findall(qn('w:gridCol'))) if tblGrid is not None else 2
    colgroup = '<colgroup>' + '<col style="min-width: 25px;">' * max(1, grid_cols) + '</colgroup>'
    return f'<table style="min-width: 50px;">{colgroup}<tbody>{"".join(rows_html)}</tbody></table>'

# --- generate HTML per bài ---
def generate_bai_html(bai_num, items):
    parts = [f'<h1>{BAI_NAMES[bai_num]}</h1>']
    pending_imgs = []
    prev_h = 1
    saw_table_since_heading = False
    just_emitted_h2 = False

    for item in items:
        if item['type'] == 'table':
            if pending_imgs:
                for b64 in pending_imgs:
                    parts.append(f'<img src="{b64}" alt="">')
                pending_imgs = []
            parts.append(table_to_html(item['table']))
            saw_table_since_heading = True
            just_emitted_h2 = False
            continue

        # --- paragraph ---
        text = item['text']
        bold = item['bold']
        imgs = item['imgs']
        indent = item['indent']
        p = item['p']

        # Empty paragraph with only images
        if not text and imgs:
            pending_imgs.extend(imgs)
            continue

        # Figure caption
        fm = RX_FIGURE.match(text)
        if fm:
            cap = escape(text)
            if pending_imgs:
                parts.append(f'<figure src="" alt="" caption="{cap}" width="" data-image-figure="" class="image-figure">')
                for b64 in pending_imgs:
                    parts.append(f'<img src="{b64}" alt="" class="resizable-img" draggable="true">')
                parts.append(f'<figcaption contenteditable="false">{cap}</figcaption></figure>')
                pending_imgs = []
            else:
                parts.append(f'<p>{cap}</p>')
            continue

        # Determine heading level
        hl = None
        p_style = item['p'].style.name if item['p'].style else ''
        if bold:
            if RX_ROMAN.match(text):
                hl = 1
                just_emitted_h2 = False
            elif RX_NOTE.match(text):
                pass
            elif RX_INSTRUCTION_START.match(text):
                pass
            elif RX_YEU_CAU.match(text):
                hl = 4
            else:
                is_buoc = bool(RX_BuOC.match(text))
                if is_buoc:
                    hl = 3 if prev_h <= 2 else min(prev_h + 1, 4)
                    just_emitted_h2 = False
                elif indent >= 3:
                    hl = 3
                    just_emitted_h2 = False
                elif p_style == 'List Paragraph':
                    if just_emitted_h2 or (prev_h >= 3 and not saw_table_since_heading):
                        hl = 3
                        just_emitted_h2 = False
                    else:
                        hl = 2
                        just_emitted_h2 = True
                elif prev_h >= 3 and not saw_table_since_heading:
                    hl = 3
                    just_emitted_h2 = False
                elif just_emitted_h2:
                    hl = 3
                    just_emitted_h2 = False
                else:
                    hl = 2
                    just_emitted_h2 = True
        elif RX_BuOC.match(text):
            hl = 3 if prev_h <= 2 else min(prev_h + 1, 4)
            just_emitted_h2 = False

        if hl is not None:
            if pending_imgs:
                for b64 in pending_imgs:
                    parts.append(f'<img src="{b64}" alt="">')
                pending_imgs = []
            parts.append(f'<h{hl}>{escape(text)}</h{hl}>')
            prev_h = hl
            saw_table_since_heading = False
        else:
            just_emitted_h2 = False
            if imgs:
                pending_imgs.extend(imgs)
            if text:
                parts.append(f'<p>{escape(text)}</p>')

    # Flush leftover images
    for b64 in pending_imgs:
        parts.append(f'<img src="{b64}" alt="">')

    return '\n'.join(parts)

# --- process ---
bai_items = {1: [], 2: [], 3: [], 4: [], 5: []}
for item in all_items:
    b = item['bai']
    if b > 0:
        bai_items[b].append(item)

conn = sqlite3.connect(DB_PATH)
cur = conn.cursor()

for b in range(1, 6):
    html = generate_bai_html(b, bai_items[b])
    note_id = NOTE_IDS[b]
    cur.execute("UPDATE notes SET content = ?, updated_at = datetime('now') WHERE id = ?",
                (html, note_id))
    print(f"Updated {BAI_NAMES[b]}: {len(html)} bytes")

conn.commit()
conn.close()
print("\nDone!")
